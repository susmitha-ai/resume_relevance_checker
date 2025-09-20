"""
Text extraction module for resumes and job descriptions.
Supports PDF and DOCX files with fallback mechanisms.
"""

import os
import io
from pathlib import Path
from typing import Optional, Union
import logging


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_file(file_input, save_extracted: bool = True) -> str:
    """
    Extract text from uploaded file (PDF or DOCX).
    
    Args:
        file_input: Streamlit file uploader object or file path
        save_extracted: Whether to save extracted text to file
    
    Returns:
        Extracted text as string
    """
    try:

        if hasattr(file_input, 'read'):
            file_content = file_input.read()
            file_name = file_input.name
            file_extension = Path(file_name).suffix.lower()
        else:

            file_path = Path(file_input)
            file_name = file_path.name
            file_extension = file_path.suffix.lower()
            with open(file_path, 'rb') as f:
                file_content = f.read()
        

        if file_extension == '.pdf':
            text = extract_from_pdf(file_content)
        elif file_extension in ['.docx', '.doc']:
            text = extract_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        

        text = clean_text(text)
        

        if save_extracted and text:
            save_path = Path("data/extracted_texts") / f"{Path(file_name).stem}.txt"
            save_path.parent.mkdir(parents=True, exist_ok=True)
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(text)
            logger.info(f"Saved extracted text to {save_path}")
        
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from {file_name}: {e}")
        raise

def extract_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF with pdfplumber fallback.
    
    Args:
        file_content: PDF file content as bytes
    
    Returns:
        Extracted text
    """
    text = ""
    

    try:
        import fitz
        doc = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text()
            text_parts.append(page_text)
        
        text = "\n".join(text_parts)
        doc.close()
        logger.info("Successfully extracted text using PyMuPDF")
        
    except Exception as e:
        logger.warning(f"PyMuPDF failed: {e}, trying pdfplumber...")
        

        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                text = "\n".join(text_parts)
            logger.info("Successfully extracted text using pdfplumber")
            
        except Exception as e:
            logger.error(f"Both PyMuPDF and pdfplumber failed: {e}")
            raise ValueError(f"Could not extract text from PDF: {e}")
    
    return text

def extract_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_content: DOCX file content as bytes
    
    Returns:
        Extracted text
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n".join(text_parts)
        logger.info("Successfully extracted text from DOCX")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting from DOCX: {e}")
        raise ValueError(f"Could not extract text from DOCX: {e}")

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw extracted text
    
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    

    import re
    text = re.sub(r'\s+', ' ', text)
    

    text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>\|\~\`]', '', text)
    

    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if line:

            if re.match(r'^\d+$', line):
                continue

            if len(line) < 3 and line.isupper():
                continue
            cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    

    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

def extract_text_from_path(file_path: Union[str, Path]) -> str:
    """
    Extract text from file path (convenience function).
    
    Args:
        file_path: Path to file
    
    Returns:
        Extracted text
    """
    return extract_text_from_file(file_path, save_extracted=True)

def get_file_info(file_input) -> dict:
    """
    Get basic information about the file.
    
    Args:
        file_input: File uploader object or file path
    
    Returns:
        Dictionary with file information
    """
    try:
        if hasattr(file_input, 'name'):
            file_name = file_input.name
            file_size = file_input.size
        else:
            file_path = Path(file_input)
            file_name = file_path.name
            file_size = file_path.stat().st_size
        
        return {
            'name': file_name,
            'size': file_size,
            'extension': Path(file_name).suffix.lower()
        }
    except Exception as e:
        logger.error(f"Error getting file info: {e}")
        return {'name': 'unknown', 'size': 0, 'extension': ''}


def test_extraction():
    """Test text extraction with sample files."""
    test_files = [
        "data/resumes/sample.pdf",
        "data/jds/sample.docx"
    ]
    
    for file_path in test_files:
        if os.path.exists(file_path):
            try:
                text = extract_text_from_path(file_path)
                print(f"Extracted {len(text)} characters from {file_path}")
                print(f"Preview: {text[:200]}...")
            except Exception as e:
                print(f"Error extracting {file_path}: {e}")
        else:
            print(f"Test file not found: {file_path}")

if __name__ == "__main__":
    test_extraction()